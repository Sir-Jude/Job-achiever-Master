from docx import Document
from docx.shared import Cm

document = Document()

# Set font
style = document.styles['Normal']
style.font.name = 'DejaVu Serif'

# Add title
document.add_heading(
    """Giulio Arpaia
    Wathever Str. 69 12045 Berlin
    willing to relocate worldwide)
    Linkedin: giulio-arpaia-5b19b8149
    Email: giulio.arpaia@gmail.com
    Mob: 0134567890            
    """, 0)

# Add paragraphs
p = document.add_paragraph()
p.add_run('Profile in comic sans').bold = True


document.save("test.docx")